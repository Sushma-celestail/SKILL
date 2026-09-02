"""
Generate me.docx — Phase 1 Complete Summary Document
Architecture Skill v1.1.0
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1F, 0x35, 0x64)   # headings
MID_BLUE    = RGBColor(0x27, 0x63, 0xAB)   # sub-headings
ACCENT      = RGBColor(0x2E, 0x86, 0xC1)   # diagram boxes / highlights
GREEN       = RGBColor(0x1E, 0x8B, 0x4C)   # pass / done
ORANGE      = RGBColor(0xD6, 0x8A, 0x00)   # pending
RED         = RGBColor(0xC0, 0x39, 0x2B)   # risk / fail
LIGHT_GRAY  = RGBColor(0xF2, 0xF2, 0xF2)
MID_GRAY    = RGBColor(0xBF, 0xBF, 0xBF)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)

# ── Helper: shade a table cell ────────────────────────────────────────────────
def shade_cell(cell, rgb: RGBColor):
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    """kwargs: top, bottom, left, right — each a dict(sz, color, val)"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            cfg = kwargs[edge]
            tag.set(qn('w:val'),   cfg.get('val',   'single'))
            tag.set(qn('w:sz'),    str(cfg.get('sz', 4)))
            tag.set(qn('w:space'), '0')
            tag.set(qn('w:color'), cfg.get('color', '000000'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)

# ── Helper: add a coloured heading paragraph ──────────────────────────────────
def heading(text, level=1, color=DARK_BLUE, space_before=18, space_after=8):
    styles = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3'}
    p = doc.add_paragraph(style=styles.get(level, 'Heading 1'))
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold       = True
    run.font.color.rgb = color
    run.font.size  = Pt(18 - (level - 1) * 3)
    return p

# ── Helper: body paragraph ────────────────────────────────────────────────────
def body(text, bold=False, italic=False, size=10.5, color=BLACK, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold            = bold
    run.italic          = italic
    run.font.size       = Pt(size)
    run.font.color.rgb  = color
    return p

# ── Helper: styled bullet ─────────────────────────────────────────────────────
def bullet(text, level=0, color=BLACK):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size      = Pt(10.5)
    run.font.color.rgb = color
    return p

# ── Helper: standard table ────────────────────────────────────────────────────
def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, DARK_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold           = True
        run.font.color.rgb = WHITE
        run.font.size      = Pt(9.5)
    # Data rows
    for ri, row in enumerate(rows):
        tr = t.add_row()
        bg = LIGHT_GRAY if ri % 2 == 0 else WHITE
        for ci, cell_text in enumerate(row):
            cell = tr.cells[ci]
            shade_cell(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
    # Column widths
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()
    return t

# ── Helper: draw a simple ASCII-art-style flow using a borderless table ───────
def flow_diagram(steps, title=""):
    """steps: list of (label, description) — renders as vertical flow boxes"""
    if title:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(f"  {title}")
        run.bold = True
        run.font.color.rgb = MID_BLUE
        run.font.size = Pt(10.5)

    for i, (label, desc) in enumerate(steps):
        # Box
        t = doc.add_table(rows=1, cols=1)
        t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = t.rows[0].cells[0]
        shade_cell(cell, ACCENT)
        cell.width = Cm(13)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label)
        r1.bold = True
        r1.font.color.rgb = WHITE
        r1.font.size = Pt(10)
        if desc:
            p.add_run("\n")
            r2 = p.add_run(desc)
            r2.font.color.rgb = RGBColor(0xD6, 0xEA, 0xFF)
            r2.font.size = Pt(9)

        # Arrow (except after last)
        if i < len(steps) - 1:
            ap = doc.add_paragraph("        ↓")
            ap.paragraph_format.space_before = Pt(0)
            ap.paragraph_format.space_after  = Pt(0)
            ap.runs[0].font.size = Pt(13)
            ap.runs[0].font.color.rgb = ACCENT
            ap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

# ── Helper: two-column side-by-side boxes ─────────────────────────────────────
def two_col_boxes(left_title, left_items, right_title, right_items,
                  left_color=GREEN, right_color=ORANGE):
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    def fill(cell, title, items, color):
        shade_cell(cell, color)
        p = cell.paragraphs[0]
        r = p.add_run(title + "\n")
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
        for item in items:
            r2 = p.add_run(f"  • {item}\n")
            r2.font.color.rgb = WHITE
            r2.font.size = Pt(9)

    fill(t.rows[0].cells[0], left_title,  left_items,  left_color)
    fill(t.rows[0].cells[1], right_title, right_items, right_color)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover_title.add_run("Architecture Skill — Phase 1 & 2")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = DARK_BLUE

cover_sub = doc.add_paragraph()
cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover_sub.add_run("Baseline Establishment")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = MID_BLUE

doc.add_paragraph()
cover_ver = doc.add_paragraph()
cover_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover_ver.add_run("Complete Summary Document")
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

meta_table = doc.add_table(rows=5, cols=2)
meta_table.style = 'Table Grid'
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ("Skill Version",    "1.1.0"),
    ("Release Date",     "2026-09-02"),
    ("Status",           "Enterprise-Governed Skill Specification"),
    ("Approver",         "Sushma S — Skill Owner / Architecture Lead"),
    ("Git Repository",   "https://github.com/Sushma-celestail/SKILL.git  |  Tag: v1.1.0"),
]
for i, (k, v) in enumerate(meta_data):
    cells = meta_table.rows[i].cells
    shade_cell(cells[0], DARK_BLUE)
    shade_cell(cells[1], LIGHT_GRAY)
    r0 = cells[0].paragraphs[0].add_run(k)
    r0.bold = True
    r0.font.color.rgb = WHITE
    r0.font.size = Pt(10)
    r1 = cells[1].paragraphs[0].add_run(v)
    r1.font.size = Pt(10)
    cells[0].width = Cm(5)
    cells[1].width = Cm(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — WHAT IS PHASE 1?
# ═══════════════════════════════════════════════════════════════════════════════

heading("1. What Is Phase 1?", 1)

body("Phase 1 — Baseline Establishment is the first maturity step in the five-phase "
     "enterprise skill improvement roadmap. Its sole purpose is control and evidence, "
     "not optimisation. Before any improvement can be verified, you need to know "
     "exactly what exists, confirm it works, record the evidence, and lock it down "
     "as the known-good reference point for all future comparisons.")

doc.add_paragraph()
body("One sentence summary:", bold=True)
body('  \u201cKnow what you have, prove it works, freeze it, and sign it off.\u201d', italic=True,
     color=MID_BLUE, size=12)

doc.add_paragraph()
heading("The Five-Phase Maturity Roadmap", 2)

phases = [
    ("Phase 1 — Baseline Audit",         "Know what you have. Prove it works. Freeze it. ← YOU ARE HERE"),
    ("Phase 2 — Context Optimisation",   "Make the skill handle large PRDs efficiently with a structured source index."),
    ("Phase 3 — Tier Routing",           "Make tier selection deterministic, repeatable, and override-tracked."),
    ("Phase 4 — Verification & Testing", "Automate artifact validation: verify.js, JSON Schema, 10 PRD fixtures."),
    ("Phase 5 — Governance & Operations","CI/CD, named owners, metrics, incident management, production approval."),
]
add_table(
    ["Phase", "Purpose"],
    phases,
    col_widths=[6, 10.5]
)

doc.add_paragraph()
heading("Why Phase 1 Comes First", 2)
body("Without a frozen, approved baseline you cannot answer the question: "
     "\u201cDid this change make things better or worse?\u201d Every future improvement "
     "is measured against Phase 1. The health checker, diff tool, and audit "
     "report exist so that question always has a verifiable answer.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — WHAT THE SKILL DOES
# ═══════════════════════════════════════════════════════════════════════════════

heading("2. What the Architecture Skill Does", 1)

body("The Architecture Skill (SKILL.md) transforms a documented Product Requirements "
     "Document (PRD) into a complete, explainable, validated software architecture. "
     "It is operated by an AI agent following the rules inside SKILL.md.")

doc.add_paragraph()

flow_diagram([
    ("Input: Documented PRD",           "Business objective, actors, functional & non-functional requirements"),
    ("Validate PRD",                     "Identify critical blockers vs non-blocking clarifications"),
    ("Extract & classify requirements",  "PRD-Stated | Architecturally-Derived | Proposed | Blocked"),
    ("Design architecture",              "Modules, Components, APIs, Workflows, Data, Security, Deployment…"),
    ("Validate & cross-check",           "§43 validation checklist + §44 consistency check + §62 automated checks"),
    ("Deliver artifacts",                "Markdown (.md) + Word (.docx) + JSON — all three every time"),
], title="How the Skill Processes a PRD")

doc.add_paragraph()
heading("What the Skill Produces", 2)

add_table(
    ["Artifact", "Format", "Purpose"],
    [
        ("Architecture document",   ".md (Markdown)",    "Human-readable, canonical source of truth"),
        ("Architecture document",   ".docx (Word)",      "Reviewer-friendly rendered copy with embedded diagrams"),
        ("Architecture data",       ".json",             "Machine-readable, downstream skill input"),
    ],
    col_widths=[5.5, 4, 7]
)

heading("What the Skill Is NOT", 2)
for item in [
    "Not a code generator",
    "Not a database schema generator",
    "Not an API implementation generator",
    "Not a tool that invents requirements — every decision traces to the PRD",
]:
    bullet(item, color=RED)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — EVERY FILE IN THE WORKSPACE
# ═══════════════════════════════════════════════════════════════════════════════

heading("3. Every File — What It Is and What It Does", 1)

body("Phase 1 produced or formalised the following files. Every file has a specific "
     "role in the governance process. Nothing is redundant.")

doc.add_paragraph()

files = [
    ("SKILL.md",
     "Canonical skill specification",
     "§1–§68",
     "The master rulebook. 68 sections defining how the AI designs architecture. "
      "This is what the AI reads and follows on every run."),
    ("skill-manifest.json",
     "Version source",
     "Phase 1",
     "Single source of truth for skill version, release date, status, and changelog. "
      "All tools read this file for version information."),
    ("CHANGELOG.md",
     "Release changelog",
     "Phase 1",
     "Auto-generated from skill-manifest.json by changelog-generate.js. "
      "Human-readable history of every version change."),
    ("references/enterprise-capability-roadmap.md",
     "Capability roadmap",
     "§57 ref",
     "Defines what each phase requires, what's implemented vs specified, "
      "completion evidence, and the enterprise readiness matrix."),
    ("governance/health-check.js",
     "Health checker",
     "Phase 1 / GAP-003",
     "Runs 27 automated checks: required files exist, manifest valid, "
      "frontmatter present, no conflict markers, all 12 enterprise sections "
      "present, SHA-256 hashes recorded. Result: HEALTHY / UNHEALTHY."),
    ("governance/change-diff.js",
     "Change diff tool",
     "Phase 1 / GAP-002",
     "Compares two SKILL.md versions. Produces section-level diff (for "
      "reviewers) and full line-level diff (for audit evidence). "
      "Run: node change-diff.js --old <old.md> --new <new.md>"),
    ("governance/changelog-generate.js",
     "Changelog generator",
     "Phase 1 / GAP-004",
     "Reads skill-manifest.json changelog array and writes CHANGELOG.md. "
      "Run: node changelog-generate.js"),
    ("governance/baselines/baseline-audit-report-v1.1.0.md",
     "Baseline audit report",
     "Phase 1 / GAP-001",
     "The signed, read-only Phase 1 evidence record. Lists all files, "
      "hashes, capabilities confirmed, gaps, risks, and approver sign-off. "
      "Approved by: Sushma S — 2026-09-02."),
    ("governance/health-reports/health-report-2026-09-02.md",
     "Health report (live run)",
     "Phase 1",
     "Auto-generated by health-check.js. Records 27 PASS, 0 FAIL, 0 WARN "
      "with full SHA-256 hashes for all three required files."),
]

add_table(
    ["File", "Role", "Linked To", "What It Does"],
    [[f[0], f[1], f[2], f[3]] for f in files],
    col_widths=[5, 3.5, 2.5, 6.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — FOLDER STRUCTURE
# ═══════════════════════════════════════════════════════════════════════════════

heading("4. Workspace Folder Structure", 1)

body("Every file has a deliberate location. The structure mirrors the intended "
     "production layout described in the enterprise-capability-roadmap.md.")

doc.add_paragraph()

# Draw structure as a styled table acting as a tree
tree_rows = [
    ("📁  Desktop/Phases/  (workspace root)",                 ""),
    ("    ├── SKILL.md",                                      "Canonical skill specification"),
    ("    ├── skill-manifest.json",                           "Version source — read by all tools"),
    ("    ├── CHANGELOG.md",                                  "Auto-generated release history"),
    ("    ├── .gitignore",                                    "Excludes node_modules, logs, OS files"),
    ("    ├── .gitattributes",                                "LF line-ending normalisation for Git"),
    ("    ├── 📁  references/",                               ""),
    ("    │   └── enterprise-capability-roadmap.md",          "Five-phase capability roadmap"),
    ("    └── 📁  governance/",                               ""),
    ("        ├── health-check.js",                           "Repository health checker (27 checks)"),
    ("        ├── change-diff.js",                            "Section + line diff between versions"),
    ("        ├── changelog-generate.js",                     "Writes CHANGELOG.md from manifest"),
    ("        ├── 📁  baselines/",                            ""),
    ("        │   └── baseline-audit-report-v1.1.0.md",      "Signed Phase 1 baseline evidence"),
    ("        └── 📁  health-reports/",                       ""),
    ("            └── health-report-2026-09-02.md",           "Live health check result (27 PASS)"),
]

t = doc.add_table(rows=len(tree_rows), cols=2)
t.style = 'Table Grid'
for i, (path_str, desc) in enumerate(tree_rows):
    cells = t.rows[i].cells
    is_folder = path_str.strip().startswith("📁") or path_str.strip().startswith("(")
    bg = RGBColor(0xE8, 0xF0, 0xFE) if is_folder else WHITE
    shade_cell(cells[0], bg)
    shade_cell(cells[1], WHITE)
    r = cells[0].paragraphs[0].add_run(path_str)
    r.font.name = "Courier New"
    r.font.size = Pt(8.5)
    r.bold = is_folder
    r2 = cells[1].paragraphs[0].add_run(desc)
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    cells[0].width = Cm(9)
    cells[1].width = Cm(8.5)

doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — THE PHASE 1 WORKFLOW
# ═══════════════════════════════════════════════════════════════════════════════

heading("5. The Phase 1 Workflow — Step by Step", 1)

body("This is the exact sequence of steps that was executed to complete Phase 1. "
     "Each step has a clear input, action, output, and the file it produced or updated.")

doc.add_paragraph()

flow_diagram([
    ("Step 1: Confirm the canonical skill file exists",
     "Input: SKILL.md (v1.1.0, 68 sections)"),
    ("Step 2: Create skill-manifest.json",
     "Version source — skill_version, release_date, status, changelog entries"),
    ("Step 3: Add enterprise-capability-roadmap.md",
     "Defines Phase 1–5 requirements and completion evidence"),
    ("Step 4: Add §68 to SKILL.md",
     "Formally specifies Phase 1 inside the skill itself"),
    ("Step 5: Build governance/health-check.js",
     "27 automated checks — files, manifest, sections, hashes"),
    ("Step 6: Build governance/change-diff.js",
     "Section-level + line-level diff between any two SKILL.md versions"),
    ("Step 7: Build governance/changelog-generate.js",
     "Reads manifest → writes CHANGELOG.md"),
    ("Step 8: Run health-check.js",
     "Result: 27 PASS, 0 FAIL → health-report-2026-09-02.md"),
    ("Step 9: Write baseline-audit-report-v1.1.0.md",
     "Files, hashes, capabilities, gaps, risks — awaiting sign-off"),
    ("Step 10: Initialise Git repo + push to GitHub",
     "git init → git add → git commit → git push → tag v1.1.0"),
    ("Step 11: Approver signs off",
     "Sushma S — Skill Owner / Architecture Lead — 2026-09-02 ✅"),
], title="Phase 1 Execution Workflow")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — HOW THE TOOLS CONNECT
# ═══════════════════════════════════════════════════════════════════════════════

heading("6. How the Tools Connect — Data Flow", 1)

body("The three governance scripts work together as a pipeline. "
     "Each reads from a source and writes to an output. "
     "No tool modifies another tool's output.")

doc.add_paragraph()

# Data flow as a table diagram
flow_rows = [
    ("skill-manifest.json",         "→",  "changelog-generate.js",        "→",  "CHANGELOG.md"),
    ("SKILL.md + skill-manifest",   "→",  "health-check.js",               "→",  "health-reports/health-report-<date>.md"),
    ("Old SKILL.md + New SKILL.md", "→",  "change-diff.js",                "→",  "change-diffs/diff-<v>-to-<v>-<date>.md"),
]

t = doc.add_table(rows=len(flow_rows) + 1, cols=5)
t.style = 'Table Grid'
headers = ["Input", "", "Tool (Script)", "", "Output"]
for ci, h in enumerate(headers):
    cell = t.rows[0].cells[ci]
    shade_cell(cell, DARK_BLUE)
    r = cell.paragraphs[0].add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(9.5)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for ri, (inp, arr1, tool, arr2, out) in enumerate(flow_rows):
    row = t.rows[ri + 1]
    bg = LIGHT_GRAY if ri % 2 == 0 else WHITE
    for ci, val in enumerate([inp, arr1, tool, arr2, out]):
        cell = row.cells[ci]
        shade_cell(cell, bg if ci != 1 and ci != 3 else WHITE)
        r = cell.paragraphs[0].add_run(val)
        r.font.size = Pt(9)
        if ci == 2:
            r.bold = True
            r.font.color.rgb = ACCENT
        if ci in (1, 3):
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r.font.size = Pt(14)
            r.font.color.rgb = MID_BLUE

col_widths = [4.5, 0.6, 4, 0.6, 4.8]
for ci, w in enumerate(col_widths):
    for row in t.rows:
        row.cells[ci].width = Cm(w)

doc.add_paragraph()

heading("When to Run Each Tool", 2)
add_table(
    ["Tool", "Run When", "Command"],
    [
        ("health-check.js",       "Before any release; after any file change",
         "node governance/health-check.js"),
        ("change-diff.js",        "Comparing two SKILL.md versions",
         "node governance/change-diff.js --old <old.md> --new SKILL.md"),
        ("changelog-generate.js", "After updating skill-manifest.json changelog",
         "node governance/changelog-generate.js"),
    ],
    col_widths=[4, 6, 7.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — HEALTH CHECK EXPLAINED
# ═══════════════════════════════════════════════════════════════════════════════

heading("7. Health Checker — 27 Checks Explained", 1)

body("health-check.js runs automatically and produces a health report. "
     "Here is what each group of checks verifies and why it matters.")

doc.add_paragraph()

check_groups = [
    ("Check Group",           "Checks",  "Why It Matters"),
    ("Required files exist",  "3",       "SKILL.md, skill-manifest.json, enterprise-capability-roadmap.md must all be present."),
    ("Manifest is valid JSON","1",       "A corrupt manifest breaks every tool that reads it."),
    ("Manifest fields present","6",      "skill_name, skill_version, release_date, status, canonical_file, changelog."),
    ("Referenced files resolve","1",     "Every file listed in manifest.references must exist on disk."),
    ("Frontmatter present",   "1",       "SKILL.md must have a valid --- delimited frontmatter block."),
    ("No conflict markers",   "1",       "<<<<<<< markers mean an unresolved merge — never release with these."),
    ("Enterprise sections",   "12",      "§57–§68 must all be present. Missing sections = incomplete enterprise spec."),
    ("SHA-256 hashes",        "3",       "Content hashes prove exactly what was in each file at baseline time."),
]
add_table(
    ["Check Group", "Checks", "Why It Matters"],
    check_groups[1:],
    col_widths=[5, 2, 10.5]
)

doc.add_paragraph()
body("Live result (2026-09-02):", bold=True)

add_table(
    ["Result", "Count"],
    [("✅  PASS", "27"), ("❌  FAIL", "0"), ("⚠️   WARN", "0"), ("Overall", "HEALTHY")],
    col_widths=[6, 11.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 — WHAT IS IMPLEMENTED VS NOT YET
# ═══════════════════════════════════════════════════════════════════════════════

heading("8. What Is Implemented vs Not Yet Implemented", 1)

body("The enterprise-capability-roadmap.md uses precise status language. "
     "Here is the honest state of every capability.")

doc.add_paragraph()

two_col_boxes(
    "✅  Implemented in Phase 1",
    [
        "Baseline-audit report (GAP-001)",
        "Change-diff tool (GAP-002)",
        "Repository health checker (GAP-003)",
        "Release-changelog generator (GAP-004)",
        "Git repo + release tag v1.1.0 (GAP-011)",
        "skill-manifest.json (version source)",
        "Approver sign-off recorded (Sushma S)",
    ],
    "⏳  Specified — Not Yet Implemented",
    [
        "Source Index Contract in §60 (Phase 2)",
        "tier_override field in §41/§61 (Phase 3)",
        "verify.js validation script (Phase 4)",
        "architecture.schema.json (Phase 4)",
        "T-001–T-010 test fixtures (Phase 4)",
        "CI/CD pipeline (Phase 5)",
        "Named Skill Owner + metrics dashboard (Phase 5)",
    ],
    left_color=RGBColor(0x1A, 0x7A, 0x40),
    right_color=RGBColor(0xB5, 0x6A, 0x00),
)

doc.add_paragraph()
heading("Gap Status Table", 2)

gaps = [
    ("GAP-001", "Baseline-audit report generator",      "✅ Implemented", "baseline-audit-report-v1.1.0.md"),
    ("GAP-002", "Change-diff comparison tool",          "✅ Implemented", "governance/change-diff.js"),
    ("GAP-003", "Repository health checker",            "✅ Implemented", "governance/health-check.js"),
    ("GAP-004", "Release-changelog generator",          "✅ Implemented", "governance/changelog-generate.js"),
    ("GAP-005", "V-001–V-010 runnable automated checks","⏳ Phase 4",     "verify.js (to build)"),
    ("GAP-006", "T-001–T-010 test fixture files",       "⏳ Phase 4",     "test-fixtures/ (to build)"),
    ("GAP-007", "Test runner",                          "⏳ Phase 4",     "runner script (to build)"),
    ("GAP-008", "JSON Schema for §46",                  "⏳ Phase 4",     "architecture.schema.json (to build)"),
    ("GAP-009", "Source Index Contract in §60",         "⏳ Phase 2",     "SKILL.md §60 update"),
    ("GAP-010", "tier_override field §41/§61",          "⏳ Phase 3",     "SKILL.md §41/§61 update"),
    ("GAP-011", "Git repository + release tag",         "✅ Implemented", "github.com/Sushma-celestail/SKILL  tag: v1.1.0"),
    ("GAP-012", "Named approver",                       "✅ Implemented", "Sushma S — 2026-09-02"),
    ("GAP-013", "CI pipeline",                          "⏳ Phase 5",     "GitHub Actions (to build)"),
    ("GAP-014", "Runtime monitoring/metrics",           "⏳ Phase 5",     "dashboard (to build)"),
]
add_table(
    ["Gap ID", "Description", "Status", "Delivered As / Next Step"],
    gaps,
    col_widths=[2, 5.5, 2.5, 7.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 — THE FIVE EVALUATION LAYERS
# ═══════════════════════════════════════════════════════════════════════════════

heading("9. The Five Evaluation Layers", 1)

body("Phases tell you what to BUILD over time. "
     "Evaluation layers tell you how to ASSESS the skill at any point. "
     "They are related but different.")

doc.add_paragraph()

add_table(
    ["Layer", "What It Evaluates", "Primary Phase"],
    [
        ("Layer 1 — Static skill review",
         "Is SKILL.md well-structured? Any contradictions? Clear outputs and boundaries?",
         "Phase 1"),
        ("Layer 2 — PRD/input evaluation",
         "Did the skill read all requirement sections? Preserve source IDs? Detect missing inputs?",
         "Phase 2"),
        ("Layer 3 — Behaviour evaluation",
         "Did it choose the right Tier? Identify real blockers? Separate PRD-Stated/Derived/Proposed/Blocked?",
         "Phase 3"),
        ("Layer 4 — Artifact verification",
         "Is JSON valid? Are IDs unique? Do Markdown, JSON, DOCX and diagrams agree?",
         "Phase 4"),
        ("Layer 5 — Human and operational review",
         "Can an architect trust and use the output? Are proposed decisions approved before implementation?",
         "Phase 5"),
    ],
    col_widths=[4, 9, 2.5]
)

doc.add_paragraph()
heading("Skill Evaluation Workflow", 2)

body("Every skill version or change passes through all five layers in sequence "
     "before it can be released.")

doc.add_paragraph()

flow_diagram([
    ("Skill version / change",
     "A new SKILL.md edit, schema change, or tool update"),
    ("Layer 1: Static skill review",
     "Frontmatter valid? References resolve? No contradictions? Enterprise sections present?"),
    ("Layer 2: Controlled PRD test fixtures",
     "Run T-001–T-010 PRDs through the skill"),
    ("Layer 3: Evaluate architecture behaviour",
     "PRD-grounded? Correct tier? Correct blockers? No invented technology?"),
    ("Layer 4: Verify generated artifacts",
     "Markdown / JSON / DOCX / diagrams / traceability all valid and consistent?"),
    ("Layer 5: Human and operational review",
     "Architecture reviewer approves. Release controls sign off."),
    ("Release skill version",
     "Tag in Git. Update skill-manifest.json. Run changelog-generate.js."),
    ("Monitor runs, defects, cost, time, rework",
     "Feeds back into next cycle"),
], title="End-to-End Evaluation and Release Workflow")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 10 — SCORECARD
# ═══════════════════════════════════════════════════════════════════════════════

heading("10. Current Maturity Scorecard", 1)

body("This is the honest assessment of where the skill stands today, "
     "after Phase 1 is complete.")

doc.add_paragraph()

scorecard = [
    ("Core PRD → Architecture design",      "9.2 / 10", "Strong and comprehensive"),
    ("PRD traceability and blocker control", "9.5 / 10", "Major strength"),
    ("Enterprise governance specification",  "8.8 / 10", "Clearly defined"),
    ("Context/token optimisation",           "7.0 / 10", "Rules defined; no indexer/telemetry yet"),
    ("Tier-routing approach",               "8.0 / 10", "Repeatable baseline, sensible expert override"),
    ("Automated verification",              "3.0 / 10", "Defined, not implemented"),
    ("Test/regression automation",          "2.5 / 10", "Test categories defined; fixtures/runner absent"),
    ("Operational release readiness",       "4.0 / 10", "Governance specified; CI, approvals, logging absent"),
]

t = doc.add_table(rows=len(scorecard) + 1, cols=3)
t.style = 'Table Grid'
for ci, h in enumerate(["Area", "Rating", "Status"]):
    cell = t.rows[0].cells[ci]
    shade_cell(cell, DARK_BLUE)
    r = cell.paragraphs[0].add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(9.5)

for ri, (area, rating, status) in enumerate(scorecard):
    row = t.rows[ri + 1]
    score = float(rating.split('/')[0].strip())
    color = GREEN if score >= 8 else (ORANGE if score >= 5 else RED)
    shade_cell(row.cells[0], LIGHT_GRAY if ri % 2 == 0 else WHITE)
    shade_cell(row.cells[1], WHITE)
    shade_cell(row.cells[2], WHITE)

    row.cells[0].paragraphs[0].add_run(area).font.size = Pt(9.5)
    r_rating = row.cells[1].paragraphs[0].add_run(rating)
    r_rating.bold = True
    r_rating.font.color.rgb = color
    r_rating.font.size = Pt(9.5)
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[2].paragraphs[0].add_run(status).font.size = Pt(9.5)

for ci, w in enumerate([7, 3, 7.5]):
    for row in t.rows:
        row.cells[ci].width = Cm(w)

doc.add_paragraph()

summary_scores = [
    ("Skill design quality",           "8.8 / 10"),
    ("Enterprise operational readiness","4.8 / 10"),
    ("Overall current maturity",       "7.0 / 10"),
]
t2 = doc.add_table(rows=len(summary_scores), cols=2)
t2.style = 'Table Grid'
for ri, (label, score) in enumerate(summary_scores):
    shade_cell(t2.rows[ri].cells[0], DARK_BLUE)
    shade_cell(t2.rows[ri].cells[1], LIGHT_GRAY)
    r1 = t2.rows[ri].cells[0].paragraphs[0].add_run(label)
    r1.bold = True
    r1.font.color.rgb = WHITE
    r1.font.size = Pt(10)
    r2 = t2.rows[ri].cells[1].paragraphs[0].add_run(score)
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = MID_BLUE
    t2.rows[ri].cells[0].width = Cm(9)
    t2.rows[ri].cells[1].width = Cm(8.5)

doc.add_paragraph()
body("The skill is an excellent enterprise specification. It is not yet a fully "
     "automated enterprise platform. Phase 4 (verify.js + JSON Schema + test fixtures) "
     "will move the automation scores sharply upward.", italic=True, color=MID_BLUE)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 11 — WHAT PHASE 1 DID NOT DO
# ═══════════════════════════════════════════════════════════════════════════════

heading("11. What Phase 1 Deliberately Did NOT Do", 1)

body("Phase 1's purpose is control and evidence — not optimisation. "
     "The following are explicitly out of scope for Phase 1 and belong to later phases.")

doc.add_paragraph()

add_table(
    ["Out of Scope for Phase 1", "Belongs To", "Why Not Phase 1"],
    [
        ("Source Index Contract (§60)",        "Phase 2", "Improving token/context efficiency is an optimisation — not a baseline concern."),
        ("tier_override field in schema",      "Phase 3", "Tier routing refinement comes after the baseline is locked."),
        ("verify.js automation script",        "Phase 4", "Automated artifact verification requires test fixtures to run against."),
        ("T-001–T-010 PRD test fixtures",      "Phase 4", "Fixtures need the verifier to be meaningful."),
        ("JSON Schema (architecture.schema)",  "Phase 4", "Schema validation is a Phase 4 verification concern."),
        ("CI/CD pipeline",                     "Phase 5", "Continuous deployment needs all prior phases working first."),
        ("Metrics dashboard",                  "Phase 5", "Nothing to monitor until the skill is running in production."),
    ],
    col_widths=[5.5, 2.5, 9.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 12 — NEXT STEPS
# ═══════════════════════════════════════════════════════════════════════════════

heading("12. Recommended Next Steps", 1)

body("Phase 1 is complete. Here is the recommended sequence for Phase 2 onwards.")

doc.add_paragraph()

add_table(
    ["Priority", "Action", "Phase", "Impact"],
    [
        ("1 — Quick win",  "Add Source Index Contract to §60 in SKILL.md",
         "Phase 2", "Lifts context/token score from 7.0. No tooling needed — spec change only."),
        ("2 — Quick win",  "Add tier_override field to §41 and §61 schemas",
         "Phase 3", "Closes the one gap in tier routing. Small SKILL.md change."),
        ("3 — Verify §62", "Confirm V-001–V-010 checks are fully specified",
         "Phase 4 prep", "Must be verified before writing verify.js."),
        ("4 — Build",      "Write verify.js + architecture.schema.json",
         "Phase 4", "Automated artifact validation. Major score improvement."),
        ("5 — Build",      "Create T-001–T-010 PRD test fixture files",
         "Phase 4", "10 controlled test PRDs covering all edge cases."),
        ("6 — Run",        "Execute verify.js against existing Service Desk output",
         "Phase 4", "First real automated verification result."),
        ("7 — Govern",     "Set up GitHub Actions CI to run health-check.js on every push",
         "Phase 5", "Baseline automation. Catches regressions instantly."),
    ],
    col_widths=[2.5, 6, 2, 7]
)

doc.add_paragraph()

heading("Phase 1 → Phase 2 Transition Criteria", 2)
body("Phase 1 is officially closed when:", bold=True)
for item in [
    "✅  baseline-audit-report-v1.1.0.md is signed off (Sushma S — 2026-09-02)",
    "✅  health-check.js runs clean (27 PASS, 0 FAIL)",
    "✅  Git tag v1.1.0 is pushed to github.com/Sushma-celestail/SKILL",
    "✅  All Phase 1 files are committed to main branch",
]:
    bullet(item, color=GREEN)

doc.add_paragraph()
body("All four criteria are now met. Phase 1 is CLOSED. ✅", bold=True,
     color=GREEN, size=12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 13 — GLOSSARY
# ═══════════════════════════════════════════════════════════════════════════════

heading("13. Glossary", 1)

glossary = [
    ("Architecture Skill",         "The AI rulebook (SKILL.md) that transforms a PRD into a validated architecture."),
    ("Baseline",                   "A frozen, approved snapshot of the skill package at a known-good state."),
    ("Canonical file",             "SKILL.md — the authoritative source. All other files derive from it."),
    ("Enterprise-governed specification", "The skill is fully specified but not yet automated/operational."),
    ("GAP-xxx",                    "A numbered gap between what is specified and what is implemented."),
    ("Health check",               "Automated run of health-check.js — 27 checks that confirm the workspace is intact."),
    ("Phase",                      "One maturity stage in the five-phase improvement roadmap (build what you need)."),
    ("Evaluation layer",           "One dimension of assessment used to score the skill (assess what you have)."),
    ("PRD",                        "Product Requirements Document — the primary input to the architecture skill."),
    ("PRD-Stated",                 "Information explicitly written in the PRD — highest evidence level."),
    ("Release tag",                "A Git tag (e.g. v1.1.0) that permanently marks a specific commit as a release."),
    ("SHA-256",                    "A cryptographic hash that uniquely identifies a file's exact content."),
    ("skill-manifest.json",        "Single source of truth for skill version, status, changelog, and file references."),
    ("Tier 1/2/3",                 "Architecture output depth levels: small module / standard app / large/complex app."),
    ("V-001–V-010",                "Ten verification checks defined in §62 — automated in Phase 4."),
    ("T-001–T-010",                "Ten test scenario categories defined in §63 — fixtures built in Phase 4."),
]

add_table(
    ["Term", "Definition"],
    glossary,
    col_widths=[5, 12.5]
)

# ═══════════════════════════════════════════════════════════════════════════════
#  PHASE 2 SECTION — SOURCE INDEXING AND CONTEXT EFFICIENCY
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading("Phase 2 — Source Indexing and Context Efficiency", 1, color=MID_BLUE)

body("Phase 2 gives the skill the ability to read large PRDs reliably, prove which "
     "PRD version was used, and detect when a source has changed after architecture "
     "work began. Without Phase 2, the skill could silently process a stale or "
     "wrong PRD and produce an architecture that traces to the wrong source.")

doc.add_paragraph()
body("One sentence summary:", bold=True)
body("  \u201cIndex the PRD before you design from it, and prove the hash matches every time.\u201d",
     italic=True, color=MID_BLUE, size=12)

doc.add_paragraph()

# ── What was already in the skill ─────────────────────────────────────────────
heading("What Was Already in the Skill (Specified)", 2)
body("These behaviours were written into SKILL.md before Phase 2 tooling existed. "
     "They were instructions the AI was expected to follow manually.")
doc.add_paragraph()
add_table(
    ["Skill Section", "Behaviour Already Specified"],
    [
        ("§52 step 6",  "Read PRD in bounded sections, not one large pass. Extract requirement-by-requirement."),
        ("§52 step 36", "Return to the literal PRD source text to verify every requirement citation before delivery."),
        ("§38",         "Calibrate output depth (Tier 1/2/3) based on PRD size and complexity."),
        ("§60",         "Use a source index for large PRDs. Persist the inventory before synthesis. Never substitute a stale source."),
        ("§58",         "Record source name, received time, and content hash. Preserve source as read-only evidence."),
    ],
    col_widths=[3, 14.5]
)

# ── What was missing ──────────────────────────────────────────────────────────
heading("What Was Missing — The Phase 2 Gap", 2)
body("The instructions existed but no tool enforced them. The AI had no way to "
     "automatically produce a structured inventory, record a hash, or detect a "
     "stale source. Every run was manual and unverifiable.")
doc.add_paragraph()
add_table(
    ["Gap ID", "Missing Capability", "Why It Matters"],
    [
        ("GAP-009", "Source Index Contract (§60 under-specified)",
         "No formal JSON schema existed. The AI could produce any format or skip the index entirely."),
        ("—",       "PRD indexer (index_prd.js)",
         "No tool parsed the PRD into a structured, locator-stamped requirement inventory."),
        ("—",       "Source hash generator",
         "No tool recorded the SHA-256 of the PRD at index time — impossible to detect tampering or drift."),
        ("—",       "Stale-source detector",
         "Nothing compared the stored hash against the current PRD before architecture synthesis began."),
        ("—",       "Attachment/stale-path checker",
         "No check that referenced attachments in the PRD actually existed on disk."),
    ],
    col_widths=[2, 5.5, 10]
)

# ── What was built ─────────────────────────────────────────────────────────────
heading("What Was Built in Phase 2", 2)
doc.add_paragraph()
p2_files = [
    ("SKILL.md §60 — Source Index Contract",
     "Spec / SKILL.md",
     "Closes GAP-009",
     "Formal JSON schemas for requirement inventory record and source hash record. "
      "Source Index Integrity Rules. Tooling reference. Now a machine-enforceable contract, not a suggestion."),
    ("phase2/index_prd.js",
     "Node.js tool",
     "Phase 2 / GAP-009",
     "PRD indexer + source hash generator. Parses any Markdown PRD section-by-section. "
      "Handles plain IDs (FR-001), escaped IDs (FR\\-001 from Word extraction), and table-cell IDs. "
      "Produces structured requirement inventory JSON + source hash record JSON."),
    ("phase2/stale-source-detect.js",
     "Node.js tool",
     "Phase 2",
     "Stale source detector. Compares current PRD SHA-256 against the hash stored in "
      "a prd-index or architecture artifact. Hard-stops with a clear error if hashes differ."),
    ("phase2/extract_prd.py",
     "Python helper",
     "Phase 2",
     "Converts PRD .docx files to Markdown using mammoth so index_prd.js can parse them. "
      "Required because the Service Desk PRD is a Word document."),
    ("phase2/PRD-service-desk-v1.md",
     "Extracted PRD",
     "Phase 2 / T-001",
     "Markdown extraction of the Internal Service Desk PRD v1. Used as the first "
      "real PRD input for the indexer. Becomes the basis for T-001 test fixture in Phase 4."),
    ("phase2/indexes/prd-index-2026-09-02.json",
     "Live index output",
     "Phase 2",
     "First real index run output. 601 requirements extracted, 597 with explicit IDs "
      "(FR-xxx, NFR-xxx, BR-xxx, OQ-xxx). 372 sections parsed. ~51,791 estimated tokens."),
    ("phase2/indexes/prd-source-hash-2026-09-02.json",
     "Source hash record",
     "Phase 2",
     "SHA-256 baseline for the Service Desk PRD. Used by stale-source-detect.js "
      "to confirm the PRD has not changed since indexing."),
]
add_table(
    ["File / Deliverable", "Type", "Linked To", "What It Does"],
    p2_files,
    col_widths=[5, 2.5, 2.5, 7.5]
)

# ── Phase 2 workflow ───────────────────────────────────────────────────────────
doc.add_page_break()
heading("Phase 2 Workflow — How It Works End to End", 2)

flow_diagram([
    ("Input: PRD file (.docx or .md)",
     "The approved, current Product Requirements Document"),
    ("Step 1: Extract to Markdown (if .docx)",
     "python phase2/extract_prd.py  →  phase2/PRD-<name>.md"),
    ("Step 2: Run PRD Indexer",
     "node phase2/index_prd.js --prd <prd.md> --out phase2/indexes"),
    ("Step 3: Indexer parses sections",
     "372 sections parsed. Bounded, section-by-section — §52 step 6 compliant"),
    ("Step 4: Indexer extracts requirements",
     "601 requirements with source_locator (section:Lnnn), type, status, module"),
    ("Step 5: Indexer records SHA-256 hash",
     "prd-source-hash-<date>.json — proof of exactly which PRD version was used"),
    ("Step 6: Run Stale Source Detector (before any architecture run)",
     "node phase2/stale-source-detect.js --prd <prd.md> --index <prd-index.json>"),
    ("Step 7: Hash check result",
     "✅ SOURCE CURRENT → proceed  |  ❌ STALE SOURCE → stop, locate correct PRD"),
    ("Step 8: Architecture synthesis uses the inventory",
     "Requirement inventory fed into §52 execution workflow. Citations verified via locators (step 36)."),
], title="Phase 2 Execution Flow")

# ── Source Index Contract explained ────────────────────────────────────────────
heading("The Source Index Contract (§60) — Key Concepts", 2)

body("Every enterprise run against a PRD with more than 10 requirements MUST "
     "produce two JSON outputs before architecture synthesis begins.")
doc.add_paragraph()

two_col_boxes(
    "Requirement Inventory Record (per requirement)",
    [
        "source_locator: section:Lnnn",
        "requirement_id: FR-001 or AUTO-xx-xxx",
        "title: text (max 120 chars)",
        "type: functional | non_functional | security…",
        "source_status: prd_stated | proposed | blocked…",
        "module: section / capability group",
        "current_or_target: current_state | target_state",
        "open_question_ids: [Q-xxx, B-xxx]",
    ],
    "Source Hash Record (per PRD file)",
    [
        "record_type: source_hash",
        "skill_version: 1.2.0",
        "prd_path: absolute or relative path",
        "prd_sha256: SHA-256 hash of the file",
        "indexed_at: ISO-8601 timestamp",
        "requirement_count: 601",
        "integrity_status: clean | warnings",
    ],
    left_color=RGBColor(0x1A, 0x5E, 0x8A),
    right_color=RGBColor(0x2E, 0x7D, 0x32),
)

# ── Integrity rules ────────────────────────────────────────────────────────────
heading("Source Index Integrity Rules", 2)
add_table(
    ["Rule", "What It Prevents"],
    [
        ("prd_sha256 MUST be checked before synthesis begins",
         "Stops architecture being generated from a PRD that changed after indexing."),
        ("Missing / moved source path = input-integrity failure",
         "Prevents silent substitution of a similarly-named but wrong document."),
        ("Every requirement MUST have a source_locator",
         "Makes §52 step 36 citation verification mechanically possible."),
        ("Inventory is read-only after creation",
         "Prevents synthesis from silently modifying the extracted requirements."),
        ("Stale attachment = warning in report + stop if critical",
         "Detects broken links to referenced files inside the PRD."),
    ],
    col_widths=[7.5, 10]
)

# ── Live results ───────────────────────────────────────────────────────────────
heading("Live Results — Service Desk PRD Index Run", 2)
add_table(
    ["Metric", "Value"],
    [
        ("PRD source file",              "PRD Internal Service Desk Tool_V1.docx"),
        ("SHA-256 (PRD)",                "8d28cdea942ebaf9655fc33f598f78e0b9c44b39b706a42e0ae4460fb477d67e"),
        ("Sections parsed",              "372"),
        ("Requirements extracted",       "601"),
        ("  — Explicit IDs (FR/NFR/BR)", "597"),
        ("  — Auto-assigned IDs",        "4"),
        ("Stale attachments",            "0"),
        ("Estimated tokens",             "~51,791"),
        ("Integrity status",             "CLEAN ✅"),
        ("Stale-source detector result", "✅ SOURCE CURRENT — Hashes match"),
    ],
    col_widths=[7, 10.5]
)

# ── What remains ──────────────────────────────────────────────────────────────
heading("What Phase 2 Did NOT Do (Deliberately)", 2)
add_table(
    ["Out of Scope for Phase 2", "Belongs To", "Why"],
    [
        ("Token/context measurement collector",  "Phase 2 (future)", "Runtime telemetry depends on the AI runtime — not available in static tools."),
        ("Automated requirement extraction validator", "Phase 4",    "Full validation needs the test fixtures and verify.js from Phase 4."),
        ("Context budget enforcement",           "Phase 4",          "Enforcement requires knowing what the runtime can handle — Phase 4 concern."),
    ],
    col_widths=[6, 3, 8.5]
)

# ── Updated gap status ─────────────────────────────────────────────────────────
heading("Updated Gap Status After Phase 2", 2)
add_table(
    ["Gap ID", "Description", "Status"],
    [
        ("GAP-009", "Source Index Contract in §60",           "✅ Closed — §60 now has formal schema + integrity rules"),
        ("GAP-001", "Baseline-audit report",                  "✅ Closed in Phase 1"),
        ("GAP-002", "Change-diff tool",                       "✅ Closed in Phase 1"),
        ("GAP-003", "Repository health checker",              "✅ Closed in Phase 1"),
        ("GAP-004", "Release-changelog generator",            "✅ Closed in Phase 1"),
        ("GAP-011", "Git repository + release tag",           "✅ Closed in Phase 1"),
        ("GAP-012", "Named approver",                         "✅ Closed in Phase 1"),
        ("GAP-005", "V-001–V-010 runnable checks",            "⏳ Phase 4"),
        ("GAP-006", "T-001–T-010 test fixture files",         "⏳ Phase 4"),
        ("GAP-007", "Test runner (verify.js)",                "⏳ Phase 4"),
        ("GAP-008", "JSON Schema for §46",                    "⏳ Phase 4"),
        ("GAP-010", "tier_override field in §41/§61",         "⏳ Phase 3 — next"),
        ("GAP-013", "CI pipeline",                            "⏳ Phase 5"),
        ("GAP-014", "Runtime monitoring/metrics",             "⏳ Phase 5"),
    ],
    col_widths=[2, 7.5, 8]
)

# ── Updated scorecard ──────────────────────────────────────────────────────────
doc.add_page_break()
heading("Updated Maturity Scorecard After Phase 2", 2)
scorecard_p2 = [
    ("Core PRD → Architecture design",      "9.2 / 10", "Unchanged — already strong"),
    ("PRD traceability and blocker control", "9.5 / 10", "Unchanged — already strong"),
    ("Enterprise governance specification",  "8.8 / 10", "Unchanged"),
    ("Context/token optimisation",           "8.0 / 10", "↑ from 7.0 — Source Index Contract + tooling implemented"),
    ("Tier-routing approach",               "8.0 / 10", "Unchanged — Phase 3 will close remaining gap"),
    ("Automated verification",              "3.0 / 10", "Unchanged — Phase 4 work"),
    ("Test/regression automation",          "2.5 / 10", "Unchanged — Phase 4 work"),
    ("Operational release readiness",       "4.0 / 10", "Unchanged — Phase 5 work"),
]
t = doc.add_table(rows=len(scorecard_p2) + 1, cols=3)
t.style = 'Table Grid'
for ci, h in enumerate(["Area", "Rating After Phase 2", "Change"]):
    cell = t.rows[0].cells[ci]
    shade_cell(cell, DARK_BLUE)
    r = cell.paragraphs[0].add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(9.5)
for ri, (area, rating, note) in enumerate(scorecard_p2):
    row = t.rows[ri + 1]
    score = float(rating.split('/')[0].strip())
    color = GREEN if score >= 8 else (ORANGE if score >= 5 else RED)
    shade_cell(row.cells[0], LIGHT_GRAY if ri % 2 == 0 else WHITE)
    shade_cell(row.cells[1], WHITE)
    shade_cell(row.cells[2], WHITE)
    row.cells[0].paragraphs[0].add_run(area).font.size = Pt(9.5)
    r_r = row.cells[1].paragraphs[0].add_run(rating)
    r_r.bold = True
    r_r.font.color.rgb = color
    r_r.font.size = Pt(9.5)
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_n = row.cells[2].paragraphs[0].add_run(note)
    r_n.font.size = Pt(9)
    r_n.font.color.rgb = GREEN if note.startswith("↑") else RGBColor(0x55, 0x55, 0x55)
for ci, w in enumerate([7, 3, 7.5]):
    for row in t.rows:
        row.cells[ci].width = Cm(w)

doc.add_paragraph()
body("Phase 2 lifted context/token optimisation from 7.0 to 8.0 by implementing the "
     "Source Index Contract, PRD indexer, and stale-source detector. "
     "The next jump will come from Phase 4 (verify.js + test fixtures).",
     italic=True, color=MID_BLUE)

# ═══════════════════════════════════════════════════════════════════════════════
#  FOOTER NOTE  (updated to v1.2.0)
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_p.add_run(
    "Architecture Skill v1.2.0  |  Phase 1 & 2 Complete  |  "
    "Approved: Sushma S \u2014 2026-09-02  |  "
    "github.com/Sushma-celestail/SKILL  tag: v1.2.0"
)
r.font.size = Pt(8.5)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r.italic = True

# ═══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════════════════════════════════════════

out_path = r"C:\Users\sushma.s\Desktop\Phases\me.docx"
doc.save(out_path)
print(f"\n✅  me.docx saved to: {out_path}\n")
